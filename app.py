# app.py — @nutritionsays · Gestión Nutricional (Ambulatorio por defecto)
# Fix: dropdown legible (menú y opciones), “Ambulatorio” como modo principal (TEE = MB×PAL [+ADE opc.]),
# “Hospitalario (avanzado)” con FE/FD/FA; form con botón Calcular; sin precálculo.

from datetime import date
from io import BytesIO
import math
import streamlit as st
import pandas as pd

# DOCX opcional
try:
    from docx import Document
    from docx.shared import Pt
    DOCX = True
except Exception:
    DOCX = False

BRAND = "@nutritionsays"
st.set_page_config(page_title=f"{BRAND} · Gestión Nutricional", page_icon="🍎", layout="centered")

# ===================== ESTILO (mejoras de legibilidad en dropdowns) =====================
st.markdown("""
<style>
/* Main claro */
.stApp, .block-container { background:#ffffff !important; color:#111 !important; }
h1,h2,h3,h4,h5, p, span, label, div, li, th, td { color:#111 !important; }

/* Sidebar oscuro legible */
section[data-testid="stSidebar"] { background:#1e1e2a !important; border-right:1px solid #141421; }
section[data-testid="stSidebar"] * { color:#f5f6fb !important; }
section[data-testid="stSidebar"] label, section[data-testid="stSidebar"] .stMarkdown p { color:#e6def7 !important; }
section[data-testid="stSidebar"] input, section[data-testid="stSidebar"] textarea {
  color:#ffffff !important; background:#111223 !important; border:1px solid #3b3b57 !important;
}
section[data-testid="stSidebar"] input::placeholder, section[data-testid="stSidebar"] textarea::placeholder { color:#cbd0ff !important; opacity:.85; }

/* Selects visibles en sidebar y en el cuerpo */
section[data-testid="stSidebar"] div[data-baseweb="select"]>div { color:#ffffff !important; }
section[data-testid="stSidebar"] div[data-baseweb="select"] svg { fill:#ffffff !important; }

/* MENÚ DESPLEGABLE (cuando abres la lista): fondo oscuro + texto blanco */
.stApp [data-baseweb="menu"] {
  background:#1b1d2c !important; color:#ffffff !important;
  border:1px solid #3b3b57 !important; box-shadow:0 4px 14px rgba(0,0,0,.35) !important;
}
.stApp [data-baseweb="menu"] * { color:#ffffff !important; }
.stApp [data-baseweb="menu"] li, .stApp [data-baseweb="menu"] div[role="option"] { color:#ffffff !important; }
.stApp [data-baseweb="menu"] li:hover, .stApp [data-baseweb="menu"] div[role="option"]:hover {
  background:#2a2e46 !important;
}

/* Tarjetas y badges */
.card { border:1px solid #e6e6ef; border-radius:14px; padding:14px; background:#fff; box-shadow:0 1px 6px rgba(0,0,0,.06); }
.kpi { font-size:1.12rem; font-weight:700; }
.badge { display:inline-block; padding:2px 8px; border-radius:999px; font-size:.82rem; font-weight:700; }
.bg-ok { background:#e7f6ef; color:#106a42; border:1px solid #c7eadb;}
.bg-warn { background:#fff3cd; color:#8a6d1a; border:1px solid #ffe29a;}
.bg-bad { background:#fde2e1; color:#a02c2a; border:1px solid #f6b1af;}
.bg-info{ background:#e7efff; color:#274c9a; border:1px solid #c6d7ff;}
@media (max-width: 480px){ .stApp { padding:.4rem; } h1{font-size:1.36rem;} h2{font-size:1.12rem;} .card{padding:10px;} }
</style>
""", unsafe_allow_html=True)

st.markdown(f"### {BRAND} · Software de Gestión Nutricional")

# ===================== Estado base =====================
if "ready" not in st.session_state: st.session_state.ready=False
if "kcal" not in st.session_state: st.session_state.kcal=None
if "payload" not in st.session_state: st.session_state.payload=None

# ===================== Catálogos =====================
EXCHANGES = {
    "Vegetales": {"kcal":25,"CHO":5,"PRO":2,"FAT":0,"portion":"1 taza crudas / 1/2 taza cocidas"},
    "Frutas": {"kcal":60,"CHO":15,"PRO":0,"FAT":0,"portion":"1 unid pequeña / 1/2 taza picada"},
    "Cereales": {"kcal":80,"CHO":15,"PRO":2,"FAT":1,"portion":"1/2 taza cocidos / 1 rebanada pan"},
    "Leguminosas": {"kcal":100,"CHO":18,"PRO":7,"FAT":1,"portion":"1/2 taza cocidas"},
    "Lácteos descremados": {"kcal":90,"CHO":12,"PRO":8,"FAT":2,"portion":"1 taza leche / yogurt natural"},
    "Proteínas magras": {"kcal":110,"CHO":0,"PRO":21,"FAT":3,"portion":"30 g cocidos"},
    "Grasas saludables": {"kcal":45,"CHO":0,"PRO":0,"FAT":5,"portion":"1 cdita (5 g)"}
}

PAL = {  # Ambulatorio (OMS/FAO aproximado)
    "Muy bajo (sedentario)":1.2, "Ligero":1.4, "Moderado":1.6, "Alto":1.75, "Muy alto":2.0
}
FA_HOSP = {"VM/Conectado":1.1,"Reposo en cama":1.2,"Deambula (ligera)":1.3}
FE = {  # estrés/enfermedad
    "Ninguno":1.0, "Cirugía menor":1.1, "Cirugía mayor":1.2, "Infección moderada":1.3,
    "Trauma huesos largos":1.25, "Politrauma":1.45, "TCE":1.6, "Quemados (≥40% SCQ)":1.8
}
FD = {"Sin FD":1.0, "Desnutrición moderada/grave (0.7)":0.7}

# ===================== Utilidades =====================
def mifflin(sex, w, h_cm, age): return 10*w + 6.25*h_cm - 5*age + (5 if sex.lower().startswith("m") else -161)
def harris_benedict(sex, w, h_cm, age):
    if sex.lower().startswith("m"): return 66.47 + (13.75*w) + (5.003*h_cm) - (6.755*age)
    return 655.09 + (9.563*w) + (1.850*h_cm) - (4.676*age)

def tee_ambulatorio(mb, pal, ade_on=False):
    base = mb * pal
    if ade_on: base *= 1.10
    return round(base)

def tee_hospitalario(mb, fa, fe, fd, ade_on=False):
    base = mb * (1.10 if ade_on else 1.0)
    return round(base * fa * fe * fd)

def kcal_target(tee, obj):
    if obj=="Pérdida de peso": return max(1000, tee - (400 if tee>=1600 else 200))
    if obj=="Ganancia (magro)": return tee + 200
    return tee

def bmi(w,hcm):
    if not w or not hcm: return None
    h=max(1e-6, hcm/100); return round(w/(h*h),2)

def whr(waist, hip): return round((waist/hip),2) if waist and hip else None
def whtr(waist, hcm): return round((waist/hcm),2) if waist and hcm else None
def homa_ir(gmgdl, ins): 
    if gmgdl and ins: return round(((gmgdl/18.0)*ins)/22.5,2)
    return None

# Durnin–Womersley + Siri (4 pliegues)
DW = {
    "F":[(17,(1.1549,0.0678)),(29,(1.1599,0.0717)),(39,(1.1423,0.0632)),(49,(1.1333,0.0612)),(120,(1.1339,0.0645))],
    "M":[(17,(1.1620,0.0630)),(29,(1.1631,0.0632)),(39,(1.1422,0.0544)),(49,(1.1620,0.0700)),(120,(1.1715,0.0779))]
}
def dw_density(sex, age, biceps, triceps, subesc, supra):
    S=max(0.1,(biceps or 0)+(triceps or 0)+(subesc or 0)+(supra or 0)); logS=math.log10(S)
    key="F" if sex.lower().startswith("f") else "M"
    coeff=None
    for up,ab in DW[key]:
        if age<=up: coeff=ab; break
    if coeff is None: coeff=DW[key][-1][1]
    a,b=coeff; return a-(b*logS)
def siri_pctfat(d): return round(((4.95/d)-4.50)*100,1)

def sodium_convert(target_mg, current_mg):
    target_mg=target_mg or 2300; current_mg=current_mg or 0
    rem=max(0,target_mg-current_mg); salt_g=round(rem/400.0,2); tsp=round(salt_g/5.0,2)
    return {"remaining_mg":rem,"salt_g":salt_g,"tsp":tsp}

def macros(kcal, pct_prot, pct_fat, pct_cho, w, pct_cho_complex=85, fat_split=(10,35,55)):
    kcal=max(0,int(kcal or 0))
    total=max(1,int(pct_prot)+int(pct_fat)+int(pct_cho))
    pct_prot=round(100*int(pct_prot)/total); pct_fat=round(100*int(pct_fat)/total); pct_cho=100-pct_prot-pct_fat
    g_prot=round((kcal*pct_prot/100)/4,1); g_fat=round((kcal*pct_fat/100)/9,1); g_cho=round((kcal*pct_cho/100)/4,1)
    gkg_prot=round(g_prot/(w or 1),2); gkg_cho=round(g_cho/(w or 1),2)
    g_cho_c=round(g_cho*(pct_cho_complex or 0)/100,1); g_cho_s=round(g_cho-g_cho_c,1)
    sat,poli,mono=fat_split; subtotal=max(1,int(sat)+int(poli)+int(mono))
    sat=pct_fat*int(sat)/subtotal; poli=pct_fat*int(poli)/subtotal; mono=pct_fat-sat-poli
    g_sat=round((kcal*sat/100)/9,1); g_poli=round((kcal*poli/100)/9,1); g_mono=round((kcal*mono/100)/9,1)
    return {"pct":{"prot":pct_prot,"fat":pct_fat,"cho":pct_cho},
            "g":{"prot":g_prot,"fat":g_fat,"cho":g_cho,"cho_c":g_cho_c,"cho_s":g_cho_s,"sat":g_sat,"poli":g_poli,"mono":g_mono},
            "gkg":{"prot":gkg_prot,"cho":gkg_cho}}

def exchanges_from_kcal(k):
    if not k or k<=0: return {g:0 for g in EXCHANGES}
    f=max(1.0, min(2.4, k/2000))
    base={"Vegetales":4,"Frutas":2,"Cereales":5,"Leguminosas":1,"Lácteos descremados":1,"Proteínas magras":4,"Grasas saludables":4}
    return {g:int(round(v*f)) for g,v in base.items()}
def distribute_by_meal(d):
    split={"Desayuno":0.25,"Merienda AM":0.10,"Almuerzo":0.30,"Merienda PM":0.10,"Cena":0.25}
    out={m:{} for m in split}
    for g,tot in d.items():
        for m,fr in split.items(): out[m][g]=round(tot*fr,1)
    return out

def ibw_hamwi(sex, height_cm):
    if not height_cm: return 0.0
    inches_over_5ft = max(0.0, (height_cm - 152.4) / 2.54)
    if sex.lower().startswith("m"):
        return 48.0 + 2.7*inches_over_5ft
    return 45.5 + 2.2*inches_over_5ft

def peso_ajustado_obesidad(actual, ibw): return ibw + 0.25*((actual or 0) - (ibw or 0)) if actual and ibw else 0.0
def ama_area(muac_cm, triceps_mm):
    if muac_cm and triceps_mm:
        tsf_cm = triceps_mm/10.0
        return round(((muac_cm - math.pi*tsf_cm)**2) / (4*math.pi), 2)
    return None

def badge(text, level="info"):
    cls = {"ok":"bg-ok","warn":"bg-warn","bad":"bg-bad","info":"bg-info"}.get(level, "bg-info")
    return f"<span class='badge {cls}'>{text}</span>"

def interp_labs(sex, labs):
    out=[]
    g=labs.get("glu",0)
    if g: out.append(("Glucosa", g, badge("Baja","warn") if g<70 else badge("Normal","ok") if g<100 else badge("Prediabetes","warn") if g<126 else badge("Diabetes","bad")))
    a1c=labs.get("a1c",0)
    if a1c: out.append(("HbA1c", a1c, badge("Normal","ok") if a1c<5.7 else badge("Prediabetes","warn") if a1c<6.5 else badge("Diabetes","bad")))
    if labs.get("homa") is not None: out.append(("HOMA-IR", labs["homa"], badge("Aceptable","ok") if labs["homa"]<2.5 else badge("↑ Resistencia","warn")))
    ldl=labs.get("ldl",0); hdl=labs.get("hdl",0); tg=labs.get("tg",0); tc=labs.get("tc",0)
    if ldl: out.append(("LDL", ldl, badge("Deseable","ok") if ldl<100 else badge("Alto","bad")))
    if hdl:
        low = 40 if sex.lower().startswith("m") else 50
        out.append(("HDL", hdl, badge("Protector","ok") if hdl>=low else badge("Bajo","bad")))
    if tg: out.append(("TG", tg, badge("Normal","ok") if tg<150 else badge("Alto","bad")))
    if tc: out.append(("CT", tc, badge("Deseable","ok") if tc<200 else badge("Alto","bad")))
    creat=labs.get("creat",0)
    if creat:
        hi=1.3 if sex.lower().startswith("m") else 1.1; lo=0.5
        out.append(("Creatinina", creat, badge("Normal","ok") if lo<=creat<=hi else badge("Alta","bad") if creat>hi else badge("Baja","warn")))
    alt=labs.get("alt",0); ast=labs.get("ast",0)
    if alt: out.append(("ALT", alt, badge("Normal","ok") if alt<=40 else badge("Alta","bad")))
    if ast: out.append(("AST", ast, badge("Normal","ok") if ast<=40 else badge("Alta","bad")))
    hb=labs.get("hb",0)
    if hb:
        lo = 13.5 if sex.lower().startswith("m") else 12.0
        hi = 17.5 if sex.lower().startswith("m") else 16.0
        out.append(("Hemoglobina", hb, badge("Normal","ok") if lo<=hb<=hi else badge("Baja","bad") if hb<lo else badge("Alta","warn")))
    ferr=labs.get("ferr",0)
    if ferr:
        lo = 24 if sex.lower().startswith("m") else 12
        hi = 336 if sex.lower().startswith("m") else 150
        out.append(("Ferritina", ferr, badge("Normal","ok") if lo<=ferr<=hi else badge("Baja","bad") if ferr<lo else badge("Alta","warn")))
    vitd=labs.get("vitd",0)
    if vitd: out.append(("Vit D", vitd, badge("Deficiencia","bad") if vitd<20 else badge("Insuficiente","warn") if vitd<30 else badge("Suficiente","ok")))
    b12=labs.get("b12",0)
    if b12: out.append(("B12", b12, badge("Baja","bad") if b12<200 else badge("Alta","warn") if b12>900 else badge("Normal","ok")))
    tsh=labs.get("tsh",0)
    if tsh: out.append(("TSH", tsh, badge("Normal","ok") if 0.4<=tsh<=4.0 else badge("Alterada","warn")))
    urea=labs.get("urea",0)
    if urea: out.append(("Urea", urea, badge("Normal","ok") if 15<=urea<=45 else badge("Alterada","warn")))
    crp=labs.get("crp",0)
    if crp: out.append(("PCR", crp, badge("Aceptable","ok") if crp<=5 else badge("Alta","bad")))
    return out

# ===================== SIDEBAR (form con botón) =====================
with st.sidebar:
    with st.form("cap"):
        st.subheader("Paciente")
        modo = st.selectbox("Modo", ["Ambulatorio (recomendado)", "Hospitalario (avanzado)"])
        nombre = st.text_input("Nombre y apellido")
        sexo = st.selectbox("Sexo biológico", ["Femenino","Masculino"])
        edad = st.number_input("Edad (años)", 1, 120, 0, step=1)
        talla_cm = st.number_input("Talla (cm)", 0, 230, 0)
        peso = st.number_input("Peso (kg)", 0.0, 300.0, 0.0, step=0.1)

        st.caption("Ecuación y factores energéticos")
        eq = st.selectbox("Ecuación de MB", ["Mifflin–St Jeor","Harris–Benedict"])
        ade_on = st.checkbox("Añadir ADE/TEF (~10%)", value=False)

        if modo.startswith("Ambulatorio"):
            pal = st.selectbox("PAL (actividad)", list(PAL.keys()), index=1)
            fa_hosp = None; fe = None; fd = None
        else:
            pal = None
            fa_hosp = st.selectbox("FA (actividad hospitalaria)", list(FA_HOSP.keys()), index=1)
            fe = st.selectbox("FE (estrés/enfermedad)", list(FE.keys()), index=0)
            fd = st.selectbox("FD (desnutrición)", list(FD.keys()), index=0)

        objetivo = st.selectbox("Objetivo", ["Pérdida de peso","Mantenimiento","Ganancia (magro)"], index=1)

        with st.expander("Antropometría (opcional)", expanded=False):
            cintura = st.number_input("Cintura (cm)", 0.0, 300.0, 0.0, step=0.1)
            cadera  = st.number_input("Cadera (cm)", 0.0, 300.0, 0.0, step=0.1)
            muac    = st.number_input("CB/MUAC (cm)", 0.0, 80.0, 0.0, step=0.1)
            p_bi = st.number_input("Bíceps (mm)", 0.0, 60.0, 0.0, step=0.5)
            p_tri = st.number_input("Tríceps (mm)", 0.0, 60.0, 0.0, step=0.5)
            p_sub = st.number_input("Subescapular (mm)", 0.0, 60.0, 0.0, step=0.5)
            p_sup = st.number_input("Suprailiaco (mm)", 0.0, 60.0, 0.0, step=0.5)
            bia_fat = st.number_input("% Grasa (BIA)", 0.0, 70.0, 0.0, step=0.1)

        with st.expander("Laboratorios (opcional)", expanded=False):
            glicemia = st.number_input("Glucosa (mg/dL)", 0.0, 800.0, 0.0, step=0.1)
            insulina = st.number_input("Insulina (µUI/mL)", 0.0, 1000.0, 0.0, step=0.1)
            hba1c = st.number_input("HbA1c (%)", 0.0, 20.0, 0.0, step=0.1)
            tc  = st.number_input("Colesterol total (mg/dL)", 0.0, 500.0, 0.0, step=0.1)
            hdl = st.number_input("HDL (mg/dL)", 0.0, 200.0, 0.0, step=0.1)
            ldl = st.number_input("LDL (mg/dL)", 0.0, 300.0, 0.0, step=0.1)
            tg  = st.number_input("Triglicéridos (mg/dL)", 0.0, 1000.0, 0.0, step=0.1)

        btn = st.form_submit_button("Calcular")

# ===================== Validación mínima =====================
if btn and peso>0 and talla_cm>0 and edad>0:
    st.session_state.ready=True
    # MB
    mb = mifflin(sexo, peso, talla_cm, edad) if eq.startswith("Mifflin") else harris_benedict(sexo, peso, talla_cm, edad)
    # TEE según modo
    if pal is not None:
        tee = tee_ambulatorio(mb, PAL[pal], ade_on)
    else:
        tee = tee_hospitalario(mb, FA_HOSP[fa_hosp], FE[fe], FD[fd], ade_on)
    kcal = kcal_target(tee, objetivo)

    st.session_state.kcal = kcal
    st.session_state.payload = {
        "nombre":nombre, "sexo":sexo, "edad":edad, "talla_cm":talla_cm, "peso":peso, "eq":eq,
        "mb":round(mb), "tee":tee, "objetivo":objetivo, "pal":pal, "fa_hosp":fa_hosp, "fe":fe, "fd":fd, "ade":ade_on,
        "ant":{"cintura":cintura,"cadera":cadera,"muac":muac,"p_bi":p_bi,"p_tri":p_tri,"p_sub":p_sub,"p_sup":p_sup,"bia_fat":bia_fat},
        "labs":{"glicemia":glicemia,"insulina":insulina,"hba1c":hba1c,"tc":tc,"hdl":hdl,"ldl":ldl,"tg":tg}
    }
elif not st.session_state.ready:
    st.warning("Completa **peso**, **talla** y **edad** y pulsa **Calcular**. No se muestran resultados hasta entonces.")
    st.stop()

# ===================== A partir de aquí todo reactivo =====================
data = st.session_state.payload
peso = data["peso"]; talla_cm = data["talla_cm"]; edad = data["edad"]; sexo=data["sexo"]
kcal = st.session_state.kcal

# Antropometría y derivados
imc = bmi(peso, talla_cm)
icc = whr(data["ant"]["cintura"], data["ant"]["cadera"])
ict = whtr(data["ant"]["cintura"], talla_cm)
homa = homa_ir(data["labs"]["glicemia"], data["labs"]["insulina"])

pct_grasa_dw=None
if sum([(data["ant"]["p_bi"] or 0),(data["ant"]["p_tri"] or 0),(data["ant"]["p_sub"] or 0),(data["ant"]["p_sup"] or 0)])>0:
    dens = dw_density(sexo, edad, data["ant"]["p_bi"], data["ant"]["p_tri"], data["ant"]["p_sub"], data["ant"]["p_sup"])
    pct_grasa_dw = siri_pctfat(dens)

# IBW / PAJ
ibw = ibw_hamwi(sexo, talla_cm)
pari = round(100*(peso/(ibw or 1)),1)
paj = peso_ajustado_obesidad(peso, ibw) if (imc and (imc>=30 or (ibw and pari>=120))) else 0.0

# ===================== Requerimientos (reactivos) =====================
st.header("Requerimientos nutricionales")
use_preset = st.checkbox("Preset rápido (Prot 20%, Grasas 30%, CHO 50%)", value=True)
if use_preset:
    pct_prot, pct_fat, sat, poli, mono, pct_cho_complex = 20, 30, 10, 35, 55, 85
else:
    c1, c2 = st.columns(2)
    with c1:
        pct_prot = st.slider("Proteínas (%)", 10, 35, 20)
        pct_fat  = st.slider("Grasas totales (%)", 20, 40, 30)
    with c2:
        sat = st.slider("De la grasa total → Saturadas (%)", 0, 15, 10)
        poli = st.slider("De la grasa total → Poliinsat. (%)", 5, 60, 35)
        mono = max(0, 100 - sat - poli)
        pct_cho_complex = st.slider("Dentro de CHO → Complejos (%)", 45, 100, 85)

pct_cho = 100 - pct_prot - pct_fat
st.info(f"CHO (%) se ajusta a: **{pct_cho}%** · Monoinsat. (%) se ajusta a: **{max(0,100-sat-poli)}%**")

mac = macros(kcal, pct_prot, pct_fat, pct_cho, peso, pct_cho_complex, fat_split=(sat, poli, max(0,100-sat-poli)))

# ===================== KPIs =====================
st.header("Resultados clínicos")
k = st.columns(3)
k[0].markdown(f"<div class='card'><div class='kpi'>IMC: {imc} kg/m²</div><div>OMS: {'Bajo peso' if imc and imc<18.5 else 'Normopeso' if imc and imc<25 else 'Sobrepeso' if imc and imc<30 else 'Obesidad I' if imc and imc<35 else 'Obesidad II' if imc and imc<40 else 'Obesidad III'}</div></div>", unsafe_allow_html=True)
# Etiqueta de TEE según modo usado
modo_txt = "MB×PAL" if data["pal"] is not None else "MB×(ADE?)×FA×FE×FD"
k[1].markdown(f"<div class='card'><div class='kpi'>MB: {data['mb']} kcal</div><div>TEE ({modo_txt}): {data['tee']} kcal</div></div>", unsafe_allow_html=True)
k[2].markdown(f"<div class='card'><div class='kpi'>Meta calórica: {kcal} kcal</div><div>kcal/kg ref.: {round(kcal/((paj or ibw or peso) or 1),2)} kcal/kg</div></div>", unsafe_allow_html=True)

k2 = st.columns(3)
k2[0].markdown(f"<div class='card'><div class='kpi'>ICC: {icc if icc is not None else '—'}</div><div>Riesgo ↑ si >0.85 (F) / >0.90 (M)</div></div>", unsafe_allow_html=True)
k2[1].markdown(f"<div class='card'><div class='kpi'>ICT: {ict if ict is not None else '—'}</div><div>Riesgo ↑ si ≥0.5</div></div>", unsafe_allow_html=True)
bf=[]
if pct_grasa_dw is not None: bf.append(f"{pct_grasa_dw}% (pliegues)")
if data["ant"]["bia_fat"]>0: bf.append(f"{data['ant']['bia_fat']}% (BIA)")
k2[2].markdown(f"<div class='card'><div class='kpi'>% Grasa: {' · '.join(bf) if bf else '—'}</div><div>Durnin–Womersley + Siri / BIA</div></div>", unsafe_allow_html=True)

# ===================== Intercambios =====================
st.header("Plan por Intercambios")
diario = exchanges_from_kcal(kcal); por_comida = distribute_by_meal(diario)
df_plan = pd.DataFrame({
    "Grupo": list(diario.keys()),
    "Raciones/día": list(diario.values()),
    "kcal/rac": [EXCHANGES[g]["kcal"] for g in diario.keys()],
    "CHO": [EXCHANGES[g]["CHO"] for g in diario.keys()],
    "PRO": [EXCHANGES[g]["PRO"] for g in diario.keys()],
    "FAT": [EXCHANGES[g]["FAT"] for g in diario.keys()],
    "Porción": [EXCHANGES[g]["portion"] for g in diario.keys()]
})
st.dataframe(df_plan, use_container_width=True, height=320)
st.dataframe(pd.DataFrame([{"Tiempo":m, **gr} for m,gr in por_comida.items()]), use_container_width=True, height=300)

# ===================== Laboratorios =====================
st.header("Laboratorios – interpretación")
lab_int = interp_labs(sexo, {"glu":data["labs"]["glicemia"],"a1c":data["labs"]["hba1c"],
                             "homa":homa,"ldl":data["labs"]["ldl"],"hdl":data["labs"]["hdl"],
                             "tg":data["labs"]["tg"],"tc":data["labs"]["tc"]})
if lab_int:
    for name, val, flag in lab_int:
        st.markdown(f"- **{name}: {val}**  {flag}", unsafe_allow_html=True)
else:
    st.info("Ingrese valores para ver interpretación.")

# ===================== Exportes (DOCX opcional) =====================
st.markdown("---")
if DOCX:
    doc = Document(); stl=doc.styles["Normal"]; stl.font.name="Calibri"; stl.font.size=Pt(11)
    doc.add_heading("Plan de alimentación – " + BRAND, 0)
    doc.add_paragraph(f"Paciente: {data['nombre'] or '—'}  |  Fecha: {date.today().isoformat()}")
    doc.add_paragraph(f"MB: {data['mb']} kcal  |  TEE: {data['tee']} kcal  |  Meta: {kcal} kcal  |  kcal/kg ref.: {round(kcal/((paj or ibw or peso) or 1),2)}")
    t=doc.add_table(rows=1, cols=7)
    for i,h in enumerate(["Lista","Raciones","kcal","CHO","PRO","FAT","Porción"]): t.rows[0].cells[i].text=h
    for g,r in diario.items():
        row=t.add_row().cells
        row[0].text=g; row[1].text=str(r); row[2].text=str(EXCHANGES[g]["kcal"])
        row[3].text=str(EXCHANGES[g]["CHO"]); row[4].text=str(EXCHANGES[g]["PRO"])
        row[5].text=str(EXCHANGES[g]["FAT"]); row[6].text=EXCHANGES[g]["portion"]
    from io import BytesIO
    bio=BytesIO(); doc.save(bio); bio.seek(0)
    st.download_button("⬇️ Descargar PLAN (DOCX)", data=bio,
                       file_name="plan_nutritionsays.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.caption("Herramienta de apoyo clínico para profesionales. Ajustar a guías y juicio clínico. © " + BRAND)
