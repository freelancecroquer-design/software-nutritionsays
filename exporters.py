import json
from datetime import date
from io import BytesIO
try:
    from docx import Document
    from docx.shared import Pt
    DOCX = True
except Exception:
    DOCX = False

def fhir_nutrition_order(payload):
    # Simplificado (válido para pruebas / PoC)
    return {
      "resourceType":"NutritionOrder",
      "status":"active",
      "intent":"order",
      "dateTime": payload["fecha"],
      "patient":{"reference": f"Patient/{payload.get('patient_id','temp')}", "display":payload["paciente"]},
      "orderer":{"display": payload["profesional"]},
      "oralDiet":{
        "type":[{"text": payload["nutrition_order"]["diet_type"]}],
        "schedule":[{"repeat":{"boundsDuration":{"value":30,"unit":"days"}}}],
        "nutrient":[
          {"modifier":{"text":"Energy"}, "amount":{"value": payload["kcal"], "unit":"kcal/d"}},
          {"modifier":{"text":"Protein"}, "amount":{"value": payload["macros"]["g"]["prot"], "unit":"g/d"}},
          {"modifier":{"text":"Fat"}, "amount":{"value": payload["macros"]["g"]["fat"], "unit":"g/d"}},
          {"modifier":{"text":"Carbohydrate"}, "amount":{"value": payload["macros"]["g"]["cho"], "unit":"g/d"}}
        ],
        "texture":[{"modifier":{"text": payload["nutrition_order"].get("texture","Normal")}}],
        "excludeFoodModifier":[{"text": e} for e in payload["nutrition_order"].get("exclusions",[])]
      },
      "supplement":[{"productName": s} for s in payload["nutrition_order"].get("supplements",[])]
    }

def fhir_nutrition_intake(payload):
    return {
      "resourceType":"NutritionIntake",
      "status":"completed",
      "occurrenceDateTime": payload["fecha"],
      "consumedItem":[
        {"type":{"text":"Menu (plan)"},
         "amount":{"value": payload["kcal"], "unit":"kcal"},
         "nutrient":[
            {"item":{"text":"Protein"}, "amount":{"value": payload["macros"]["g"]["prot"], "unit":"g"}},
            {"item":{"text":"Fat"}, "amount":{"value": payload["macros"]["g"]["fat"], "unit":"g"}},
            {"item":{"text":"Carbohydrate"}, "amount":{"value": payload["macros"]["g"]["cho"], "unit":"g"}}
         ]}
      ],
      "subject":{"display": payload["paciente"]},
      "recorded":{"value": payload["fecha"]}
    }

def build_docx_note(kind, payload):
    if not DOCX: return None
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(11)
    doc.add_heading(f"HISTORIA CLÍNICA NUTRICIONAL – {kind.upper()}", level=1)
    doc.add_paragraph(f"Fecha: {payload['fecha']}   Profesional: {payload['profesional']}   Paciente: {payload['paciente']}")
    doc.add_heading("Evaluación (A)", level=2)
    doc.add_paragraph(payload["evaluation"])
    doc.add_heading("Diagnóstico (D)", level=2)
    for pes in payload["pes_list"]:
        doc.add_paragraph(f"- {pes}")
    doc.add_heading("Intervención (I)", level=2)
    doc.add_paragraph(payload["prescription"])
    doc.add_heading("Monitoreo/Evaluación (ME)", level=2)
    doc.add_paragraph(payload["monitoring"])
    # Requerimientos
    doc.add_heading("Requerimientos", level=2)
    m = payload["macros"]
    doc.add_paragraph(f"Energía: {payload['kcal']} kcal/d  ({payload['kcal_kg']} kcal/kg)")
    doc.add_paragraph(f"Proteínas: {m['pct']['prot']}% → {m['g']['prot']} g ({payload['gkg_prot']} g/kg)")
    doc.add_paragraph(f"Grasas: {m['pct']['fat']}% → {m['g']['fat']} g (Sat {m['g']['sat']} g, Poli {m['g']['poli']} g, Mono {m['g']['mono']} g)")
    doc.add_paragraph(f"CHO: {m['pct']['cho']}% → {m['g']['cho']} g (Complejos {m['g']['cho_c']} g, Simples {m['g']['cho_s']} g)")
    # Sodio
    s = payload["sodium"]
    doc.add_paragraph(f"Sodio objetivo: {s['target_mg']} mg; Consumido: {s['current_mg']} mg; Remanente: {s['remaining_mg']} mg")
    doc.add_paragraph(f"≈ {s['salt_g']} g NaCl ( {s['tsp']} cdtas )")
    bio = BytesIO(); doc.save(bio); bio.seek(0); return bio
